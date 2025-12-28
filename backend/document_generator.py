from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, register_element_cls
from docx.oxml.ns import nsdecls
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
from io import BytesIO
import os
from datetime import datetime
from hijri_converter import Gregorian

# Source: https://stackoverflow.com/a/70598444
# Author: skyway
# License: CC BY-SA 4.0
# Floating image implementation for python-docx

class CT_Anchor(BaseOxmlElement):
    """
    ``<w:anchor>`` element, container for a floating image.
    """
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        """
        Return a new ``<wp:anchor>`` element populated with the values passed
        as parameters.
        """
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = 'Picture %d' % shape_id
        anchor.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        """
        Return a new `wp:anchor` element containing the `pic:pic` element
        specified by the argument values.
        """
        pic_id = 0  # Word doesn't seem to use this, but does not omit it
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
            '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
            '           %s>\n'
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionH>\n'
            '  <wp:positionV relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionV>\n'
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:wrapNone/>\n'
            '  <wp:docPr id="666" name="unnamed"/>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:anchor>' % ( nsdecls('wp', 'a', 'pic', 'r'), int(pos_x), int(pos_y) )
        )


def new_pic_anchor(part, image_descriptor, width, height, pos_x, pos_y):
    """Return a newly-created `w:anchor` element.

    The element contains the image specified by *image_descriptor* and is scaled
    based on the values of *width* and *height*.
    """
    rId, image = part.get_or_add_image(image_descriptor)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = part.next_id, image.filename
    return CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, pos_x, pos_y)


def add_float_picture(p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0):
    """Add float picture at fixed position `pos_x` and `pos_y` to the top-left point of page.
    """
    run = p.add_run()
    anchor = new_pic_anchor(run.part, image_path_or_stream, width, height, pos_x, pos_y)
    run._r.add_drawing(anchor)

# Register the element
register_element_cls('wp:anchor', CT_Anchor)

def create_vacation_form(context):
    """
    Generates a vacation form from a template and returns it as a byte stream.

    Args:
        context (dict): A dictionary containing the data to be rendered in the template.

    Returns:
        BytesIO: A memory stream containing the generated DOCX document.
    """
    doc = DocxTemplate("backend/templates/vacation_template.docx")

    # Extract signature paths before rendering
    employee_sig_path = context.pop('employee_signature_path', None)
    manager_sig_path = context.pop('manager_signature_path', None)

    # Set empty strings for signature placeholders in template
    context['employee_signature'] = ''
    context['manager_signature'] = ''

    doc.render(context)

    # Save to BytesIO first
    temp_stream = BytesIO()
    doc.save(temp_stream)
    temp_stream.seek(0)

    # Re-open with python-docx to add floating signatures
    document = Document(temp_stream)

    # Add signatures as floating images if they exist
    # Position in Inches: measured from top-left corner of page
    if employee_sig_path and os.path.exists(employee_sig_path):
        # Employee signature: horizontal 1.8", vertical 4"
        p = document.paragraphs[0] if document.paragraphs else document.add_paragraph()
        add_float_picture(p, employee_sig_path, width=Inches(1.1),
                         pos_x=Inches(1.8), pos_y=Inches(4))

    if manager_sig_path and os.path.exists(manager_sig_path):
        # Manager signature: horizontal 1.18", vertical 6.4"
        p = document.paragraphs[0] if document.paragraphs else document.add_paragraph()
        add_float_picture(p, manager_sig_path, width=Inches(1.1),
                         pos_x=Inches(1.18), pos_y=Inches(6.4))

    # Save final document
    final_stream = BytesIO()
    document.save(final_stream)
    final_stream.seek(0)

    return final_stream

# Report translations
REPORT_TRANSLATIONS = {
    'en': {
        'title': 'Employee Dashboard Report',
        'report_period': 'Report Period',
        'period': 'Period',
        'date_range': 'Date Range',
        'to': 'to',
        'employee_profile': 'Employee Profile',
        'name': 'Name',
        'position': 'Position',
        'email': 'Email',
        'unit': 'Unit',
        'total_vacation_balance': 'Total Vacation Balance',
        'earned': 'Earned',
        'used': 'Used',
        'available': 'Available',
        'period_statistics': 'Period Statistics',
        'leaves_taken_in_period': 'Leaves Taken in Period',
        'requests_in_period': 'Requests in Period',
        'days': 'days',
        'contract_status': 'Contract Status',
        'contract_end_date': 'Contract End Date',
        'days_remaining': 'Days Remaining',
        'expiring_soon': '(Expiring Soon)',
        'todays_attendance': "Today's Attendance",
        'status': 'Status',
        'leave_requests_in_period': 'Leave Requests in Period',
        'recent_leave_requests': 'Recent Leave Requests',
        'type': 'Type',
        'start_date': 'Start Date',
        'end_date': 'End Date',
        'duration': 'Duration',
        'no_leave_requests': 'No leave requests in this period.',
        'team_overview': 'Team Overview',
        'team_summary': 'Team Summary:',
        'total_members': 'Total Members',
        'currently_on_leave': 'Currently On Leave',
        'total_leaves_taken_in_period': 'Total Leaves Taken in Period',
        'average_available_balance': 'Average Available Balance',
        'employee_name': 'Employee Name',
        'available_balance': 'Available Balance',
        'leaves_in_period': 'Leaves in Period',
        'current_status': 'Current Status',
        'leave_types_breakdown': 'Leave Types Breakdown',
        'annual': 'Annual Leave',
        'sick': 'Sick Leave',
        'emergency': 'Emergency Leave',
        'Present': 'Present',
        'On Leave': 'On Leave',
        'Approved': 'Approved',
        'Rejected': 'Rejected',
        'Pending': 'Pending',
        'Cancelled': 'Cancelled',
        'ytd': 'YEAR TO DATE',
        'last_30': 'LAST 30 DAYS',
        'last_60': 'LAST 60 DAYS',
        'last_90': 'LAST 90 DAYS',
        'full_year': 'FULL CONTRACT YEAR',
        'custom': 'CUSTOM RANGE'
    },
    'ar': {
        'title': 'تقرير لوحة القيادة للموظف',
        'report_period': 'فترة التقرير',
        'period': 'الفترة',
        'date_range': 'النطاق الزمني',
        'to': 'إلى',
        'employee_profile': 'الملف الشخصي للموظف',
        'name': 'الاسم',
        'position': 'المسمى الوظيفي',
        'email': 'البريد الإلكتروني',
        'unit': 'الوحدة',
        'total_vacation_balance': 'رصيد الإجازات الكلي',
        'earned': 'المستحق',
        'used': 'المستخدم',
        'available': 'المتوفر',
        'period_statistics': 'إحصائيات الفترة',
        'leaves_taken_in_period': 'الإجازات المأخوذة في الفترة',
        'requests_in_period': 'الطلبات في الفترة',
        'days': 'أيام',
        'contract_status': 'حالة العقد',
        'contract_end_date': 'تاريخ انتهاء العقد',
        'days_remaining': 'الأيام المتبقية',
        'expiring_soon': '(قرب الانتهاء)',
        'todays_attendance': 'حضور اليوم',
        'status': 'الحالة',
        'leave_requests_in_period': 'طلبات الإجازة في الفترة',
        'recent_leave_requests': 'طلبات الإجازة الأخيرة',
        'type': 'النوع',
        'start_date': 'تاريخ البداية',
        'end_date': 'تاريخ النهاية',
        'duration': 'المدة',
        'no_leave_requests': 'لا توجد طلبات إجازة في هذه الفترة.',
        'team_overview': 'نظرة عامة على الفريق',
        'team_summary': 'ملخص الفريق:',
        'total_members': 'إجمالي الأعضاء',
        'currently_on_leave': 'حالياً في إجازة',
        'total_leaves_taken_in_period': 'إجمالي الإجازات المأخوذة في الفترة',
        'average_available_balance': 'متوسط الرصيد المتوفر',
        'employee_name': 'اسم الموظف',
        'available_balance': 'الرصيد المتوفر',
        'leaves_in_period': 'الإجازات في الفترة',
        'current_status': 'الحالة الحالية',
        'leave_types_breakdown': 'تفصيل أنواع الإجازات',
        'annual': 'إجازة اعتيادية',
        'sick': 'إجازة مرضية',
        'emergency': 'إجازة طارئة',
        'Present': 'حاضر',
        'On Leave': 'في إجازة',
        'Approved': 'معتمد',
        'Rejected': 'مرفوض',
        'Pending': 'قيد الانتظار',
        'Cancelled': 'ملغي',
        'ytd': 'من بداية السنة حتى الآن',
        'last_30': 'آخر 30 يوماً',
        'last_60': 'آخر 60 يوماً',
        'last_90': 'آخر 90 يوماً',
        'full_year': 'سنة العقد الكاملة',
        'custom': 'فترة مخصصة'
    }
}

def format_date_for_report(date_str, date_system='gregorian', language='en'):
    """
    Format a date string according to the specified calendar system and language.

    Args:
        date_str: Date string in YYYY-MM-DD format
        date_system: 'gregorian' or 'hijri'
        language: 'en' or 'ar'

    Returns:
        Formatted date string
    """
    try:
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")

        if date_system == 'hijri':
            # Convert to Hijri
            hijri_date = Gregorian(date_obj.year, date_obj.month, date_obj.day).to_hijri()

            hijri_months_ar = [
                'محرم', 'صفر', 'ربيع الأول', 'ربيع الآخر', 'جمادى الأولى', 'جمادى الآخرة',
                'رجب', 'شعبان', 'رمضان', 'شوال', 'ذو القعدة', 'ذو الحجة'
            ]
            hijri_months_en = [
                'Muharram', 'Safar', 'Rabi\' al-Awwal', 'Rabi\' al-Thani',
                'Jumada al-Awwal', 'Jumada al-Thani', 'Rajab', 'Sha\'ban',
                'Ramadan', 'Shawwal', 'Dhul-Qi\'dah', 'Dhul-Hijjah'
            ]

            month_name = hijri_months_ar[hijri_date.month - 1] if language == 'ar' else hijri_months_en[hijri_date.month - 1]

            if language == 'ar':
                return f"{hijri_date.day} {month_name} {hijri_date.year} هـ"
            else:
                return f"{month_name} {hijri_date.day}, {hijri_date.year} AH"
        else:
            # Gregorian
            if language == 'ar':
                locale = 'ar-SA'
            else:
                locale = 'en-US'

            # Format using month name
            month_names_ar = [
                'يناير', 'فبراير', 'مارس', 'أبريل', 'مايو', 'يونيو',
                'يوليو', 'أغسطس', 'سبتمبر', 'أكتوبر', 'نوفمبر', 'ديسمبر'
            ]
            month_names_en = [
                'January', 'February', 'March', 'April', 'May', 'June',
                'July', 'August', 'September', 'October', 'November', 'December'
            ]

            month_name = month_names_ar[date_obj.month - 1] if language == 'ar' else month_names_en[date_obj.month - 1]

            if language == 'ar':
                return f"{date_obj.day} {month_name} {date_obj.year} م"
            else:
                return f"{month_name} {date_obj.day}, {date_obj.year}"
    except Exception as e:
        return date_str

def create_dashboard_report(data):
    """
    Generates a comprehensive dashboard report with full language and calendar support.
    data: dict containing user, balance, contract, attendance info, period filters, language, and date_system.
    """
    # Get language and date system preferences
    language = data.get('language', 'en')
    date_system = data.get('date_system', 'gregorian')
    t = REPORT_TRANSLATIONS[language]

    doc = Document()

    # Title
    heading = doc.add_heading(t['title'], 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Period Information (if provided)
    if 'filter_type' in data and 'period_start' in data and 'period_end' in data:
        doc.add_heading(t['report_period'], level=1)
        p = doc.add_paragraph()
        filter_type_key = data['filter_type']
        filter_type_display = t.get(filter_type_key, filter_type_key.replace('_', ' ').upper())
        p.add_run(f"{t['period']}: ").bold = True
        p.add_run(f"{filter_type_display}\n")
        p.add_run(f"{t['date_range']}: ").bold = True
        period_start_formatted = format_date_for_report(data['period_start'], date_system, language)
        period_end_formatted = format_date_for_report(data['period_end'], date_system, language)
        p.add_run(f"{period_start_formatted} {t['to']} {period_end_formatted}")

    # Employee Info
    doc.add_heading(t['employee_profile'], level=1)
    p = doc.add_paragraph()
    p.add_run(f"{t['name']}: ").bold = True
    p.add_run(f"{data['name_ar'] if language == 'ar' else data['name_en']}\n")
    p.add_run(f"{t['position']}: ").bold = True
    p.add_run(f"{data['position_ar'] if language == 'ar' else data['position_en']}\n")
    p.add_run(f"{t['email']}: ").bold = True
    p.add_run(f"{data['email']}\n")
    p.add_run(f"{t['unit']}: ").bold = True
    p.add_run(f"{data['unit_ar'] if language == 'ar' else data['unit_en']}")

    # Vacation Balance (Total)
    doc.add_heading(t['total_vacation_balance'], level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = t['earned']
    hdr_cells[1].text = t['used']
    hdr_cells[2].text = t['available']

    row_cells = table.add_row().cells
    row_cells[0].text = str(data['balance_earned'])
    row_cells[1].text = str(data['balance_used'])
    row_cells[2].text = str(data['balance_available'])

    # Period Statistics (if provided)
    if 'period_leaves_taken' in data:
        doc.add_heading(t['period_statistics'], level=1)
        p = doc.add_paragraph()
        p.add_run(f"{t['leaves_taken_in_period']}: ").bold = True
        p.add_run(f"{data['period_leaves_taken']} {t['days']}\n")
        p.add_run(f"{t['requests_in_period']}: ").bold = True
        p.add_run(f"{data['period_requests_count']}")

    # Contract Status
    doc.add_heading(t['contract_status'], level=1)
    p = doc.add_paragraph()
    p.add_run(f"{t['contract_end_date']}: ").bold = True
    contract_end_formatted = format_date_for_report(data['contract_end_date'], date_system, language) if data['contract_end_date'] != 'N/A' else data['contract_end_date']
    p.add_run(f"{contract_end_formatted}\n")
    p.add_run(f"{t['days_remaining']}: ").bold = True

    days_run = p.add_run(f"{data['days_remaining']}")
    if data['days_remaining'] <= 40:
        days_run.font.color.rgb = RGBColor(255, 0, 0)
        p.add_run(f" {t['expiring_soon']}").font.color.rgb = RGBColor(255, 0, 0)

    # Attendance
    doc.add_heading(t['todays_attendance'], level=1)
    p = doc.add_paragraph()
    p.add_run(f"{t['status']}: ").bold = True
    status_text = t.get(data['attendance_status'], data['attendance_status'])
    p.add_run(f"{status_text}")

    # Requests in Period
    requests_heading = t['leave_requests_in_period'] if 'filter_type' in data else t['recent_leave_requests']
    doc.add_heading(requests_heading, level=1)
    if data['requests']:
        req_table = doc.add_table(rows=1, cols=5)
        req_table.style = 'Table Grid'
        hdr = req_table.rows[0].cells
        hdr[0].text = t['type']
        hdr[1].text = t['start_date']
        hdr[2].text = t['end_date']
        hdr[3].text = t['duration']
        hdr[4].text = t['status']

        for req in data['requests']:
            row = req_table.add_row().cells
            row[0].text = t.get(req['vacation_type'].lower(), req['vacation_type'])
            row[1].text = format_date_for_report(req['start_date'], date_system, language)
            row[2].text = format_date_for_report(req['end_date'], date_system, language)
            row[3].text = f"{req['duration']} {t['days']}"
            row[4].text = t.get(req['status'], req['status'])
    else:
        doc.add_paragraph(t['no_leave_requests'])

    # Team Overview (For Managers/Admins)
    if 'team_data' in data and data['team_data']:
        doc.add_page_break()
        doc.add_heading(t['team_overview'], level=1)

        # Summary statistics
        if data['team_data']:
            total_team_leaves = sum(m.get('total_leaves_taken', 0) for m in data['team_data'])
            avg_balance = sum(m.get('vacation_balance', 0) for m in data['team_data']) / len(data['team_data'])
            on_leave_count = sum(1 for m in data['team_data'] if m.get('current_status') == 'On Leave')

            p = doc.add_paragraph()
            p.add_run(f"{t['team_summary']}\n").bold = True
            p.add_run(f"{t['total_members']}: {len(data['team_data'])}\n")
            p.add_run(f"{t['currently_on_leave']}: {on_leave_count}\n")
            if 'filter_type' in data:
                p.add_run(f"{t['total_leaves_taken_in_period']}: {total_team_leaves} {t['days']}\n")
            p.add_run(f"{t['average_available_balance']}: {avg_balance:.1f} {t['days']}")

            doc.add_paragraph()  # Spacing

        # Individual team members table
        team_table = doc.add_table(rows=1, cols=5)
        team_table.style = 'Table Grid'
        hdr = team_table.rows[0].cells
        hdr[0].text = t['employee_name']
        hdr[1].text = t['position']
        hdr[2].text = t['available_balance']
        hdr[3].text = t['leaves_in_period']
        hdr[4].text = t['current_status']

        for member in data['team_data']:
            row = team_table.add_row().cells
            member_name = member['name_ar'] if language == 'ar' else member['name_en']
            member_position = member.get('position_ar' if language == 'ar' else 'position_en', 'N/A')
            row[0].text = member_name
            row[1].text = member_position
            row[2].text = f"{member.get('vacation_balance', 0)} {t['days']}"
            row[3].text = f"{member.get('total_leaves_taken', 0)} {t['days']}"
            row[4].text = t.get(member.get('current_status', 'Present'), member.get('current_status', 'Present'))

        # Leave details breakdown with dates (if available)
        has_leave_details = any(member.get('leaves_details') for member in data['team_data'])
        if has_leave_details:
            doc.add_paragraph()
            doc.add_heading(t['leave_types_breakdown'], level=2)
            for member in data['team_data']:
                if member.get('leaves_details'):
                    member_name = member['name_ar'] if language == 'ar' else member['name_en']
                    p = doc.add_paragraph()
                    p.add_run(f"{member_name}:").bold = True

                    # Create a table for this member's leave details
                    leave_table = doc.add_table(rows=1, cols=4)
                    leave_table.style = 'Table Grid'
                    leave_hdr = leave_table.rows[0].cells
                    leave_hdr[0].text = t['type']
                    leave_hdr[1].text = t['start_date']
                    leave_hdr[2].text = t['end_date']
                    leave_hdr[3].text = t['duration']

                    for leave in member['leaves_details']:
                        leave_row = leave_table.add_row().cells
                        leave_row[0].text = t.get(leave['type'].lower(), leave['type'])
                        leave_row[1].text = format_date_for_report(leave['start_date'], date_system, language)
                        leave_row[2].text = format_date_for_report(leave['end_date'], date_system, language)
                        leave_row[3].text = f"{leave['duration']} {t['days']}"

                    doc.add_paragraph()  # Spacing between members

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

